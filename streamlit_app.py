import streamlit as st
import pandas as pd
import numpy as np
from sklearn.metrics.pairwise import cosine_similarity
from sklearn.feature_extraction.text import TfidfVectorizer
import os
import textract
from io import StringIO
from PyPDF2 import PdfReader
from zipfile import BadZipFile
from sqlalchemy import create_engine, Column, String, Float, Integer
from sqlalchemy.ext.declarative import declarative_base
from sqlalchemy.orm import sessionmaker
import xml.etree.ElementTree as ET
from docx import Document as DocxDocument

# Define a directory to store existing documents
DATA_DIRECTORY = "existing_doc"

# Ensure the data directory exists
if not os.path.exists(DATA_DIRECTORY):
    os.makedirs(DATA_DIRECTORY)

# Set up database
DATABASE_URL = "sqlite:///inventory.db"
engine = create_engine(DATABASE_URL)
Base = declarative_base()
Session = sessionmaker(bind=engine)
session = Session()

# Define Document model
class Document(Base):
    __tablename__ = 'documents'
    id = Column(Integer, primary_key=True, autoincrement=True)
    filename = Column(String, unique=True, nullable=False)
    content = Column(String, nullable=False)
    keywords = Column(String, nullable=True)

# Create the table
Base.metadata.create_all(engine)

# Create UI Pages
st.sidebar.title("Navigation")
page = st.sidebar.radio("Go to", ["Upload Document", "Data Analysis Result"])

def load_text_from_file(uploaded_file):
    file_extension = uploaded_file.name.split('.')[-1].lower()
    if file_extension == 'csv':
        return pd.read_csv(uploaded_file)
    elif file_extension == 'xlsx':
        try:
            return pd.read_excel(uploaded_file, engine='openpyxl')
        except BadZipFile:
            st.error("Error: The uploaded Excel file appears to be corrupted or not a valid Excel file.")
            return None
    elif file_extension == 'txt':
        return uploaded_file.read().decode('utf-8')
    elif file_extension == 'pdf':
        try:
            reader = PdfReader(uploaded_file)
            text = ""
            for page in reader.pages:
                text += page.extract_text() or ""
            return text
        except Exception as e:
            st.error(f"Error reading PDF file: {e}")
            return None
    elif file_extension in ['doc', 'docx']:
        try:
            doc = DocxDocument(uploaded_file)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            return text
        except Exception as e:
            st.error(f"Error reading Word document: {e}")
            return None
    elif file_extension == 'datadesign':
        try:
            tree = ET.parse(uploaded_file)
            root = tree.getroot()
            return ET.tostring(root, encoding='unicode')
        except Exception as e:
            st.error(f"Error reading DATADESIGN (XML) file: {e}")
            return None
    else:
        st.error("Unsupported file type.")
        return None

if page == "Upload Document":
    # Page 1: Upload Document
    st.title('Upload Document')

    # File uploader to upload a new document
    uploaded_file = st.file_uploader("Upload your file", type=['csv', 'xlsx', 'txt', 'pdf', 'doc', 'docx', 'datadesign'])
    if uploaded_file is not None:
        # Read the uploaded file content
        file_content = load_text_from_file(uploaded_file)
        if isinstance(file_content, pd.DataFrame):
            st.write("Uploaded Document Preview:")
            st.write(file_content.head())
            # Convert DataFrame to CSV format string
            file_content_str = file_content.to_csv(index=False)
        elif isinstance(file_content, str) and file_content:
            st.write("Uploaded Document Preview:")
            st.write(file_content[:500])  # Display the first 500 characters
            file_content_str = file_content
        else:
            file_content_str = None

        if file_content_str:
            # Save the uploaded file to the existing_documents folder
            file_path = os.path.join(DATA_DIRECTORY, uploaded_file.name)
            with open(file_path, 'wb') as f:
                f.write(uploaded_file.getbuffer())

            # Save the uploaded file content to the database
            new_document = Document(filename=uploaded_file.name, content=file_content_str)
            session.add(new_document)
            try:
                session.commit()
                st.success(f"File '{uploaded_file.name}' uploaded and saved successfully.")
            except Exception as e:
                session.rollback()
                st.error(f"Error saving file to the database: {e}")

    # List all existing documents from the database
    st.header("Existing Documents")
    documents = session.query(Document).all()
    if documents:
        for doc in documents:
            with st.expander(f"{doc.filename}"):
                st.write(doc.content[:500])  # Display the first 500 characters of the content
                delete_button = st.button(f"Delete {doc.filename}", key=f"delete_{doc.id}")
                if delete_button:
                    # Delete document from the database and file system
                    session.delete(doc)
                    try:
                        session.commit()
                        # Remove the file from the existing_documents folder
                        os.remove(os.path.join(DATA_DIRECTORY, doc.filename))
                        st.success(f"Document '{doc.filename}' deleted successfully.")
                    except Exception as e:
                        session.rollback()
                        st.error(f"Error deleting document from the database: {e}")
    else:
        st.write("No existing documents available.")

elif page == "Data Analysis Result":
    # Page 2: Data Analysis Result
    st.title('Data Analysis Result')
    uploaded_file = st.file_uploader("Upload a Document to Compare", type=['csv', 'xlsx', 'txt', 'pdf', 'doc', 'docx', 'datadesign'])
    if uploaded_file is not None:
        # Read the uploaded document into a DataFrame or text
        uploaded_df_or_text = load_text_from_file(uploaded_file)
        if isinstance(uploaded_df_or_text, str):
            uploaded_text = uploaded_df_or_text
            st.write("Uploaded Document Preview:")
            st.write(uploaded_text[:500])  # Display the first 500 characters

            # Load all existing documents from the database for comparison
            existing_documents = session.query(Document).all()
            document_names = [doc.filename for doc in existing_documents]
            document_contents = [doc.content for doc in existing_documents if doc.content]

            # Calculate similarity between the uploaded document and existing documents using TF-IDF
            if document_contents:
                corpus = [uploaded_text] + document_contents
                vectorizer = TfidfVectorizer(max_features=500, min_df=0.1)
                tfidf_matrix = vectorizer.fit_transform(corpus)

                # Compute cosine similarity between the uploaded document and existing documents
                similarity_matrix = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:]).flatten()

                # Find the most similar document
                most_similar_index = np.argmax(similarity_matrix)
                most_similar_document = document_names[most_similar_index]
                similarity_score = similarity_matrix[most_similar_index]

                st.markdown(f"<h3 style='color:blue; font-weight:bold;'>The most similar existing document is: '{most_similar_document}'</h3>", unsafe_allow_html=True)
                with open(os.path.join(DATA_DIRECTORY, most_similar_document), 'rb') as f:
                    st.download_button(label='Download Similar Document', data=f, file_name=most_similar_document)
                st.markdown(f"<h3 style='color:green; font-weight:bold;'>Similarity Score: {similarity_score:.2f}</h3>", unsafe_allow_html=True)
            else:
                st.warning("No existing documents found for comparison.")

        elif isinstance(uploaded_df_or_text, pd.DataFrame):
            # Convert the DataFrame to a CSV string
            uploaded_text = uploaded_df_or_text.to_csv(index=False)
            st.write("Uploaded Document Preview:")
            st.write(uploaded_df_or_text.head())

            # Load all existing documents from the database for comparison
            existing_documents = session.query(Document).all()
            document_names = [doc.filename for doc in existing_documents]
            document_contents = [doc.content for doc in existing_documents if doc.content]

            # Calculate similarity between the uploaded document and existing documents using TF-IDF
            if document_contents:
                corpus = [uploaded_text] + document_contents
                vectorizer = TfidfVectorizer()
                tfidf_matrix = vectorizer.fit_transform(corpus)

                # Compute cosine similarity between the uploaded document and existing documents
                similarity_matrix = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:]).flatten()

